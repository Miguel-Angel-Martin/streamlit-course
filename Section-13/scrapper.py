import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import re

def scrapper():
    st.subheader("Wikipedia web scrapper")    
    search_term = st.text_input("", placeholder="Search in Wikipedia", autocomplete="off")
    language = "es"
    headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) "
                  "Chrome/120.0.0.0 Safari/537.36"
}
    if search_term:
        url = f"https://{language}.wikipedia.org/wiki/{search_term.replace(' ', '_')}"
        st.write(f"URL: {url}")
    
        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            st.write(f"Status Code: {response.status_code}")
            soup = BeautifulSoup(response.text, "html.parser")
            title = soup.find("h1").text
            st.write(title)
            #paragraph= soup.find("p").text
            #st.write(paragraph)
            word_text=[]
            paragraphs = soup.find_all("p")
            for paragraph in paragraphs[:10]:
                expresion= re.sub(r'\[\d+]', '', paragraph.text)
                word_text.append(expresion)
                st.write(expresion)
            if st.button("Generate Docx", type="primary"):
                doc= Document()
                doc.add_heading(title, level=1)
                for p in word_text:
                    doc.add_paragraph(p)
                buffer= BytesIO()
                doc.save(buffer)
                st.download_button(
                    label="Download Docx",
                    data=buffer,
                    file_name=f"{title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except requests.exceptions.RequestException as e:
            st.error(f"Error fetching URL: {e}")
        except Exception as e:
            st.error(f"Error: {e}")
        

